# encoding:utf-8
import os
import traceback
from datetime import datetime

import docxtpl
import docx
from django.http import StreamingHttpResponse, HttpResponse
from django.shortcuts import render, redirect
from django.utils.encoding import escape_uri_path
from docx import Document
from docx.shared import Cm
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, InlineImage
from Register.models import User
from docx.shared import Mm
import time


# Create your views here.
def third(request):
    ctx = {}
    status = request.COOKIES.get('is_login')
    if not status:
        return redirect('/login/')
    return render(request, 'third.html', ctx)


def responseFile(request):
    username = request.COOKIES.get('username')
    password = request.COOKIES.get('password')
    list = User.objects.filter(username=username, password=password).first()

    projectName = request.COOKIES.get('projectName')
    projectNumber = request.COOKIES.get('projectNumber')
    SupplierName = list.SupplierName
    bossName = request.COOKIES.get('bossName')
    phone = list.phone
    address = list.address
    email = list.email
    sdk = list.SupplierDepositBank
    sdan = list.SupplierCorporateAccountNumber
    # 法定代表人名称
    fddbrmc = request.COOKIES.get('fddbrmc')
    # 被授权人姓名及身份证代码
    xmjsfzdm = request.COOKIES.get('xmjsfzdm')
    # 被授权人电话
    bsqrdh = request.COOKIES.get('bsqrdh')
    pay = request.COOKIES.get('pay')
    wordPay = request.COOKIES.get('wordPay')
    tm = request.COOKIES.get('time')
    hangye = request.COOKIES.get('hy')
    timemore = request.COOKIES.get('timemore')
    cyry = list.NumberOfEmployees
    zcze = list.TotalAssets
    bzrq = request.COOKIES.get('bzrq')


    yysr = list.AnnualOperatingIncome
    isxxqy = list.IsMicroEnterprise

    gongzhang = r"./img/{}_gz.png".format(username)
    qianming = r"./img/{}_sf.png".format(username)
    beishouqvan = r"./img/{}_bsqr.png".format(username)
    if isxxqy:
        xxqy = "小型企业"
    else:
        xxqy = "大型企业"
    print(phone)

    # 日期
    year = datetime.now().year
    month = datetime.now().month
    day = datetime.now().day
    if bzrq == "":
        bzrq = "{}年{}月{}日".format(year, month, day)
    document = Document(r"./statics/docx/temp.docx")
    replace_dict = {
        # 项目名称
        "purchaseDemandName": projectName,
        # 项目编号
        "businessId": projectNumber,
        # 供应商名称
        "UserName": SupplierName,
        # 联系电话
        "personTel": phone,
        # 联系地址
        "lxdz": address,
        # 电子函件/邮箱
        "email": email,
        # 日期
        "year": year,
        "month": month,
        "day": day,
        # 采购人名称
        "purchasing": bossName,
        # 供销商开户银行
        "qybank": sdk,
        # 账号
        "qyzh": sdan,
        # 法定代表人名称
        'fddbrmc': fddbrmc,
        # 被授权人姓名及身份证代码
        'xmjsfzdm': xmjsfzdm,
        # 被授权人电话
        'bsqrdh': bsqrdh,
        'pay': pay,
        'wordPay': wordPay,
        'time': tm,
        'hangye': hangye,
        'cyry': cyry,
        'yysr': yysr,
        # 小型企业
        "qylx": xxqy,
        # 日期
        "bztime": "{}年{}月{}日".format(year, month, day),
        # 资产总额
        "zcze": zcze,
        "bzrq": bzrq,
        "gq": timemore
     }

    document = check_and_change(document, replace_dict)
    filename = r"./statics/user/{}responseFile.docx".format(username)
    document.save(filename)
    solve_docx1(filename, filename, gongzhang, qianming)

    source_file_path_list = [filename]
    document = Document(r"./statics/docx/temp_end.docx")

    temp_end = r"./statics/user/{}_temp_end.docx".format(username)
    document = check_and_change(document, replace_dict)
    document.save(temp_end)

    solve_docx_end(temp_end, temp_end, gongzhang, qianming, beishouqvan)

    document = Document(r"./statics/docx/temp_zxqy.docx")

    temp_zxqy = r"./statics/user/{}_temp_zxqy.docx".format(username)
    document = check_and_change(document, replace_dict)
    document.save(temp_zxqy)

    solve_docx_zxqy(temp_zxqy, temp_zxqy, gongzhang)

    dt = request.COOKIES.get('detail')
    dv = request.COOKIES.get('deviate')
    ip = request.COOKIES.get('impl')
    if dt is not None:
        source_file_path_list.append(r"./tmp/{}".format(dt))
    source_file_path_list.append(temp_end)
    if dv is not None:
        source_file_path_list.append(r"./tmp/{}".format(dv))
    if ip is not None:
        source_file_path_list.append(r"./tmp/{}".format(ip))
    source_file_path_list.append(temp_zxqy)
    final_path = r"./tmp/{}_final.docx".format(username)
    merge_doc(source_file_path_list, final_path)
    print(final_path)

    def down_chunk_file_manager(file_path, chuck_size=1024):
        with open(file_path, "rb") as file:
            while True:
                chuck_stream = file.read(chuck_size)
                if chuck_stream:
                    yield chuck_stream
                else:
                    break

    dir_path = './img/username/'

    fm = [final_path]
    pic_list = get_img_file(dir_path)
    docx_path = r"./statics/docx/fj.docx"
    target_path = r"./tmp/{}_fj.docx".format(username)

    if pic_list != []:
        path_other = r"./tmp/{}_other.docx".format(username)
        add_other(username, path_other)
        fm.append(path_other)
    fm.append(target_path)

    if not add_f(username, docx_path, target_path):
        return render(request, 'second.html')

    merge_doc(fm, final_path)

    response = StreamingHttpResponse(down_chunk_file_manager(final_path))
    response['Content-Type'] = 'application/octet-stream'
    the_file_name = "响应文件.docx"
    # response["Content-Disposition"] = "attachment; filename*=UTF-8''{}".format(escape_uri_path(filename))
    response["Content-Disposition"] = "attachment; filename*=UTF-8''{}".format(escape_uri_path(the_file_name))

    print(escape_uri_path(the_file_name))

    return response

def solve_docx1(source_path, target_path, gongzhang, qianming):
    tpl = DocxTemplate(source_path)
    touming = r"./img/touming.png"
    if tpl:
        try:
            if not os.path.exists(gongzhang):
                gongzhang = touming
            tpl.replace_pic("Picture 1", gongzhang)
            # tpl.replace_pic("图片 6", gongzhang)
            # tpl.replace_pic("图片 5", gongzhang)
            # tpl.replace_pic("图片 2", gongzhang)

            if not os.path.exists(qianming):
                qianming = touming
            tpl.replace_pic("图片 3", qianming)
            tpl.replace_pic("图片 2", qianming)
            tpl.save(target_path)

        except Exception as e:
            print("=======================================")
            print(traceback.format_exc())

def solve_docx_end(source_path, target_path, gongzhang, qianming, beishouqvan):
    tpl = DocxTemplate(source_path)
    touming = r"./img/touming.png"
    if tpl:
        try:
            if not os.path.exists(gongzhang):
                gongzhang = touming
            tpl.replace_pic("Picture 1", gongzhang)
            tpl.replace_pic("Picture 3", gongzhang)
            tpl.replace_pic("Picture 4", gongzhang)

            if not os.path.exists(qianming):
                qianming = touming
            tpl.replace_pic("图片 8", qianming)
            tpl.replace_pic("图片 4", qianming)

            if not os.path.exists(beishouqvan):
                beishouqvan = touming
            tpl.replace_pic("图片 2", beishouqvan)

            tpl.save(target_path)
        except Exception as e:
            print(traceback.format_exc())
def solve_docx_zxqy(source_path, target_path, gongzhang):
    tpl = DocxTemplate(source_path)
    touming = r"./img/touming.png"
    if tpl:
        try:
            if not os.path.exists(gongzhang):
                gongzhang = touming
            tpl.replace_pic("Picture 1", gongzhang)

            tpl.save(target_path)
        except Exception as e:
            print(traceback.format_exc())

def zxqy(request):
    username = request.COOKIES.get('username')
    password = request.COOKIES.get('password')
    list = User.objects.filter(username=username, password=password).first()

    projectName = request.COOKIES.get('projectName')
    projectNumber = request.COOKIES.get('projectNumber')
    SupplierName = list.SupplierName
    bossName = request.COOKIES.get('bossName')
    phone = list.phone
    address = list.address
    email = list.email
    sdk = list.SupplierDepositBank
    sdan = list.SupplierCorporateAccountNumber
    zcze = list.TotalAssets
    cyry = list.NumberOfEmployees
    yysr = list.AnnualOperatingIncome
    isxxqy = list.IsMicroEnterprise
    bzrq = request.COOKIES.get('bzrq')
    # 日期
    year = datetime.now().year
    month = datetime.now().month
    day = datetime.now().day
    if bzrq == "":
        bzrq = "{}年{}月{}日".format(year, month, day)
    if isxxqy:
        xxqy = "小型企业"
    else:
        xxqy = "大型企业"
    hangye = request.COOKIES.get('hy')
    pay = request.COOKIES.get('pay')
    wordPay = request.COOKIES.get('wordPay')

    print(phone)

    # 日期
    year = datetime.now().year
    month = datetime.now().month
    day = datetime.now().day
    if bzrq == "":
        bzrq = "{}年{}月{}日".format(year, month, day)
    document = Document(r"./statics/docx/zxqy.docx")
    replace_dict = {
        # 项目名称
        "purchaseDemandName": projectName,
        # 项目编号
        "businessId": projectNumber,
        # 供应商名称
        "UserName": SupplierName,
        # 联系电话
        "personTel": phone,
        # 联系地址
        "lxdz": address,
        # 电子函件/邮箱
        "email": email,
        # 日期
        "year": year,
        "month": month,
        "day": day,
        # 采购人名称
        "purchasing": bossName,
        # 供销商开户银行
        "qybank": sdk,
        # 账号
        "qyzh": sdan,
        # 日期
        "bztime": "{}年{}月{}日".format(year, month, day),
        # 资产总额
        "zcze": zcze,
        # 从业人员
        "cyry": cyry,
        # 营业收入
        "yysr": yysr,
        # 小型企业
        "qylx": xxqy,
        # 所属行业
        "hangye": hangye,
        'pay': pay,
        'wordPay': wordPay,
        "bzrq": bzrq,
    }

    document = check_and_change(document, replace_dict)
    filename = r"./tmp/{}_zxqy.docx".format(username)
    document.save(filename)

    rp = fr"./img/{username}_gz.png"
    solve_docx_zxqy(filename, filename, rp)
    # if not replace_zxqy(filename, rp):
    #     return render(request, 'second.html')

    def down_chunk_file_manager(file_path, chuck_size=1024):
        with open(file_path, "rb") as file:
            while True:
                chuck_stream = file.read(chuck_size)
                if chuck_stream:
                    yield chuck_stream
                else:
                    break

    response = StreamingHttpResponse(down_chunk_file_manager(filename))
    response['Content-Type'] = 'application/octet-stream'
    the_file_path = "中小企业声明函.docx"
    response["Content-Disposition"] = "attachment; filename*=UTF-8''{}".format(escape_uri_path(the_file_path))
    # response["Content-Disposition"] = "attachment; filename={}".format("生成文件.docx")
    print(escape_uri_path(the_file_path))

    return response


def check_and_change(document, replace_dict):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
   （key 和 value 都是replace_dict中的键值对。）
    """
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    # print(str(key) + "->" + str(value))
                    # print(str(value))
                    para.runs[i].text = para.runs[i].text.replace(str(key), str(value))
    return document


def file_download(request):
    # do something...
    with open('file_name.txt') as f:
        c = f.read()
    return HttpResponse(c)


# def merge_doc(source_file_path_list, target_file_path):
#     '''
#     合并多个docx文件
#     :param source_file_path_list: 源文件路径列表
#     :param target_file_path: 目标文件路径
#     :return:
#     '''
#     # 填充分页符号文档
#     page_break_doc = Document()
#     page_break_doc.add_page_break()
#     # 定义新文档
#     target_doc = Document(source_file_path_list[0])
#     target_composer = Composer(target_doc)
#     for i in range(len(source_file_path_list)):
#
#         # 跳过第一个作为模板的文件
#         if i == 0:
#             continue
#         # 填充分页符文档
#         target_composer.append(page_break_doc)
#         # 拼接文档内容
#         f = source_file_path_list[i]
#         # =======================================
#         target_composer.append(Document(f))
#         # =======================================
#     # 保存目标文档
#     target_composer.save(target_file_path)

def merge_doc(source_file_path_list, target_file_path):
    new_document = Document()
    composer = Composer(new_document)
    page_break_doc = Document()
    for fn in source_file_path_list:
        composer.append(Document(fn))
        composer.append(page_break_doc)
    composer.save(target_file_path)


def replace_picture(final_path, replace_img_path):
    tpl = DocxTemplate(final_path)
    try:
        if tpl:
            try:
                tpl.replace_pic("Picture 3", replace_img_path)
            except:
                print("3不存在")
            try:
                tpl.replace_pic("Picture 6", replace_img_path)
            except:
                print("6不存在")
            try:
                # tpl.replace_pic("Picture 9", replace_img_path)
                tpl.replace_pic("图片 9", replace_img_path)
            except:
                print("9不存在")
            try:
                tpl.replace_pic("图片 11", replace_img_path)
            except:
                print("11不存在")
            try:
                tpl.replace_pic("图片 12", replace_img_path)
            except:
                print("12不存在")
        tpl.save(final_path)
        return True
    except Exception as e:
        print(traceback.format_exc())
        return False


def replace_sf(final_path, replace_img_path):
    tpl = DocxTemplate(final_path)
    try:
        if tpl:
            # tpl.replace_pic("Picture 1", replace_img_path)
            tpl.replace_pic("图片 2", replace_img_path)
        tpl.save(final_path)
        return True
    except Exception as e:
        print(traceback.format_exc())
        return False


def replace_zxqy(final_path, replace_img_path):
    tpl = DocxTemplate(final_path)
    try:
        if tpl:
            tpl.replace_pic("Picture 2", replace_img_path)
        tpl.save(final_path)
        return True
    except Exception as e:
        print(traceback.format_exc())
        return False

def add_f(username, docx_path, target_path):
    try:
        fzm = r"./img/{}_fzm.png".format(username)
        fbm = r"./img/{}_fbm.png".format(username)
        bzm = r"./img/{}_bzm.png".format(username)
        bbm = r"./img/{}_bbm.png".format(username)
        yyzz = r"./img/{}_yyzz.png".format(username)

        # 创建docx对象
        daily_docx = docxtpl.DocxTemplate(docx_path)

        # 创建2张图片对象
        if os.path.exists(fzm):
            insert_image1 = docxtpl.InlineImage(daily_docx, fzm, width=Mm(140))
        else:
            insert_image1 = ''

        if os.path.exists(fbm):
            insert_image2 = docxtpl.InlineImage(daily_docx, fbm, width=Mm(140))
        else:
            insert_image2 = ''

        if os.path.exists(bzm):
            insert_image3 = docxtpl.InlineImage(daily_docx, bzm, width=Mm(140))
        else:
            insert_image3 = ''

        if os.path.exists(bbm):
            insert_image4 = docxtpl.InlineImage(daily_docx, bbm, width=Mm(140))
        else:
            insert_image4 = ''

        if os.path.exists(yyzz):
            insert_image5 = docxtpl.InlineImage(daily_docx, yyzz, width=Mm(140))
        else:
            insert_image5 = ''

        # 渲染内容
        context = {
            "fzm": insert_image1,
            "fbm": insert_image2,
            "bzm": insert_image3,
            "bbm": insert_image4,
            "yyzz": insert_image5
        }

        # 渲染docx
        daily_docx.render(context)
        # 保存docx
        daily_docx.save(target_path)
        return True
    except Exception as e:
        print(traceback.format_exc())
        return False

def get_img_file(file_name):
    imagelist = []
    for parent, dirnames, filenames in os.walk(file_name):
        for filename in filenames:
            if filename.lower().endswith(
                    ('.bmp', '.dib', '.png', '.jpg', '.jpeg', '.pbm', '.pgm', '.ppm', '.tif', '.tiff')):
                imagelist.append(os.path.join(parent, filename))
        return imagelist
def add_other(username, final_path):
    tpl = DocxTemplate(r'./statics/docx/qita.docx')

    dir_path = './img/username/'

    pic_list = get_img_file(dir_path)
    print(pic_list)
    imgs = []

    for i in pic_list:
        imgs.append(InlineImage(tpl, i, height=Mm(100), width=Mm(100)))
    context = {
        'images': imgs
    }
    tpl.render(context)
    tpl.save(final_path)